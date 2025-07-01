# llm_utils.py
# Este m√≥dulo utiliza la API de OpenAI para generar descripciones, keywords
# y estimaciones de tono emocional a partir de conceptos o publicaciones.

import logging
import os
from dotenv import load_dotenv  # Carga variables de entorno desde un archivo .env
from openai import OpenAI, RateLimitError  # Cliente oficial para la API de OpenAI
from openai._exceptions import APIConnectionError  # Agrega esta l√≠nea para importar APIConnectionError
import ast  # Permite evaluar strings como estructuras de Python de forma segura
import re
from app.models.publicacion import Publicacion
from app.models.keyword import Keyword
from app.mongo.mongo_keywords import create_keyword
import math
import json
from datetime import datetime
from typing import List
from app.mongo.mongo_fuentes import get_fuentes_dict
from io import BytesIO
from docx import Document
from app.models.publicacion import Publicacion
from bson import ObjectId

from app.mongo.mongo_conceptos import get_collection as get_conceptos_collection


def get_openai_client():
    global open_ai_client
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key or not api_key.startswith("sk-"):
        raise Exception("‚ùå Falta la variable OPENAI_API_KEY en el entorno.")
    model = os.getenv("OPENAI_MODEL", "gpt-3.5-turbo")  # Modelo de OpenAI a utilizar

    # en caso que este configurado por limite el uso de chatgpt, se puede usar open router
    if "true" == os.getenv("USE_OPEN_ROUTER", "false").lower():
        # Si se usa open router, se debe usar el cliente de open router
        logging.info("Usando OpenRouter como cliente de OpenAI.")
        open_ai_client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=os.getenv("OPEN_ROUTER_API_KEY"),
        )
        model = f"openai/{model}"  # Cambia el modelo para que sea compatible con open router
    else:
        open_ai_client = OpenAI(api_key=api_key)
    return open_ai_client, model


def get_gpt_response(messages, temperature):
    """
    Env√≠a un mensaje a la API de OpenAI y devuelve la respuesta generada.
    :param messages:
    :param temperature:
    :return:
    """

    # Cargar la API Key desde el archivo .env
    client, model = get_openai_client()
    response = None
    try:
        response = client.chat.completions.create(model=model, messages=messages, temperature=temperature)
        response = response.choices[0].message.content.strip()
    except APIConnectionError as e:
        logging.error(f"Error al conectar con la API de OpenAi: \n{e}")
        raise e

    except Exception as e:
        logging.error(f"Error al generar la descripci√≥n del concepto: \n{e}")
        raise e
    except RateLimitError as e:
        logging.error("üö´ L√≠mite de cuota alcanzado. Verifica tu facturaci√≥n en OpenAI.\n La IA no puede generar una respuesta en este momento: se ha alcanzado el l√≠mite de uso.")
        raise e
    # Se devuelve el contenido limpio del mensaje generado
    return response if response else "Error al obtener respuesta de la IA."

# ------------------------------------------------------------------------------
# Genera una descripci√≥n clara y concisa de un concepto dado
def generar_descripcion_concepto(nombre_concepto: str) -> str:
    """
    Genera una descripci√≥n de 4 frases sobre el concepto dado.
    """
    prompt = (
        f"Redacta un p√°rrafo de 4 frases explicando de manera clara y espec√≠fica el tema: '{nombre_concepto}'. Desde un punto de vista de la actualidad y las noticias publicadas"
        "Usa lenguaje t√©cnico pero comprensible. No repitas palabras innecesarias. No uses comillas ni comillas dobles."
    )
    messages = [
        {"role": "system", "content": "Eres un redactor experto en comunicaci√≥n clara y precisa."},
        {"role": "user", "content": prompt}
    ]

    response = get_gpt_response(messages, 0.7)
    return response if response else "Error al generar la descripci√≥n."


# ------------------------------------------------------------------------------
# Extrae y guarda 10 keywords representativas a partir de una descripci√≥n
def generar_keywords_descriptivos(descripcion_concepto: str) -> list[Keyword]:
    """
    Genera 10 keywords a partir de la descripci√≥n del concepto,
    las guarda en Mongo (si no existen) y devuelve objetos Keyword con _id.
    """
    prompt = (
        f"A partir de esta descripci√≥n:\n\n\"{descripcion_concepto}\"\n\n"
        "Devuelve un array de 10 keywords representativas del tema. "
        "Cada keyword puede tener hasta 4 palabras. Deben ser √∫tiles para an√°lisis sem√°ntico con IA. "
        "Devuelve solo una lista de Python v√°lida, sin explicaci√≥n adicional."
    )
    messages = [
        {"role": "system", "content": "Eres un analista experto en inteligencia sem√°ntica y taxonom√≠as."},
        {"role": "user", "content": prompt}
    ]

    try:
        response = get_gpt_response(messages, 0.5)
        if not response:
            raise ValueError("Respuesta vac√≠a del modelo.")

        def limpiar_respuesta_de_codigo(respuesta: str) -> str:
            if respuesta.startswith("```") and respuesta.endswith("```"):
                lineas = respuesta.strip().splitlines()
                return "\n".join(lineas[1:-1]).strip()
            return respuesta.strip()

        raw_keywords = ast.literal_eval(limpiar_respuesta_de_codigo(response))

        if not isinstance(raw_keywords, list) or not all(isinstance(k, str) for k in raw_keywords):
            raise TypeError("La respuesta no es una lista de strings v√°lida.")

        keywords_guardadas = []

        for nombre in raw_keywords:
            nombre_limpio = nombre.strip()
            keyword_obj = Keyword(nombre=nombre_limpio)

            # Guardar usando create_keyword
            respuesta, status = create_keyword(keyword_obj)

            if status in (200, 201):  # ya existe o se acaba de crear
                json_data = respuesta.get_json()
                keyword_obj._id = json_data.get("_id")
                keywords_guardadas.append(keyword_obj)
            else:
                logging.warning(f"No se pudo guardar la keyword: {nombre_limpio} (status {status})")

        return keywords_guardadas

    except (ValueError, SyntaxError, TypeError) as e:
        logging.error(f"‚ùå Error al generar o guardar keywords: {e}\nRespuesta recibida: {response}")
        raise e

# ------------------------------------------------------------------------------
# Estima el tono emocional del t√≠tulo de una publicaci√≥n entre 1 (muy negativo) y 10 (muy positivo)
def estimar_tono_publicacion(publicacion) -> int:
    """
    Usa el t√≠tulo de una publicaci√≥n para estimar su tono del 1 (muy negativo) al 10 (muy positivo).
    Si hay enlaces en el t√≠tulo, los elimina antes del an√°lisis.
    """

    # Elimina enlaces (http, https o www) del t√≠tulo
    titulo_limpio = re.sub(r'https?://\S+|www\.\S+', '', publicacion.titulo).strip()

    prompt = (
        f"A partir del siguiente t√≠tulo:\n\n\"{titulo_limpio}\"\n\n"
        "Valora el tono emocional del contenido impl√≠cito en el t√≠tulo, del 1 al 10, "
        "donde 1 es muy negativo, 5 es neutro, y 10 es muy positivo. "
        "Devuelve solo el n√∫mero entero, sin explicaci√≥n adicional."
    )
    messages = [
        {"role": "system", "content": "Eres un analista experto en comunicaci√≥n y an√°lisis emocional de lenguaje."},
        {"role": "user", "content": prompt}
    ]

    # Se env√≠a el prompt al modelo para que devuelva un n√∫mero emocional
    tono_str = get_gpt_response(messages, 0)
    try:
        return int(tono_str)  # Convertimos el resultado a entero
    except ValueError:
        raise ValueError(f"Respuesta inesperada del modelo: {tono_str}")  # Manejo de errores si no devuelve un n√∫mero

# Resume el contenido de una publicaci√≥n, reformulando con sin√≥nimos para evitar infracci√≥n de copyright
def resumir_contenido_reformulado(publicacion: Publicacion, keywords_dict=None, max_tokens=600) -> Publicacion:
    """
    Resume y reformula el contenido de una publicaci√≥n utilizando un LLM.
    Se incorporan las keywords relacionadas (por ID) si est√°n disponibles.
    El resumen evita frases literales y debe mantenerse entre 6 y 10 l√≠neas.
    """
    if not publicacion.contenido.strip():
        raise ValueError("El contenido est√° vac√≠o o no disponible.")

    if len(publicacion.contenido) > 15000:
        publicacion.contenido = publicacion.contenido[:15000]

    # Extraer nombres de keywords si est√°n disponibles
    keyword_nombres = []
    if hasattr(publicacion, "keywords_relacionadas_ids") and publicacion.keywords_relacionadas_ids:
        if keywords_dict is not None:
            keyword_nombres = [keywords_dict.get(str(kid), "") for kid in publicacion.keywords_relacionadas_ids]
        else:
            # Si no se pasa el diccionario, obtenerlas directamente (esto requiere acceso a Mongo)
            from app.mongo.mongo_keywords import get_keywords_by_ids
            keywords = get_keywords_by_ids(publicacion.keywords_relacionadas_ids)
            keyword_nombres = [kw["nombre"] for kw in keywords if "nombre" in kw]

    prompt = (
        "Resume el siguiente art√≠culo en un m√≠nimo de 6 y un m√°ximo de 10 l√≠neas. "
        "Evita repetir frases textuales del texto original: reformula con sin√≥nimos y un estilo claro. "
        "Enf√≥cate especialmente en los temas relacionados con las siguientes palabras clave:\n\n"
        f"{', '.join(keyword_nombres)}\n\n"
        "El tono debe ser profesional y directo. Aqu√≠ est√° el texto completo del art√≠culo:\n\n"
        f"{publicacion.contenido}"
    )

    messages = [
        {"role": "system", "content": "Eres un asistente experto en resumir y reformular art√≠culos de prensa evitando plagio."},
        {"role": "user", "content": prompt}
    ]

    resumen = get_gpt_response(messages, temperature=0.7)
    publicacion.contenido = resumen
    return publicacion



def analizar_publicacion(publicacion, max_tokens=600):
    """
    Analiza una publicaci√≥n:
    - Resume y reformula el contenido.
    - Estima el tono emocional del t√≠tulo (1-10).
    
    Devuelve el objeto Publicacion con:
    - contenido: actualizado con el resumen.
    - tono: nuevo atributo con el valor del tono emocional.
    """
    if not publicacion.contenido.strip():
        raise ValueError("El contenido est√° vac√≠o o no disponible.")

    if len(publicacion.contenido) > 25000:
        publicacion.contenido = publicacion.contenido[:25000]

    titulo_limpio = re.sub(r'https?://\S+|www\.\S+', '', publicacion.titulo).strip()

    prompt = (
    f"T√≠tulo: \"{titulo_limpio}\"\n\n"
    f"Contenido:\n{publicacion.contenido}\n\n"
    "Primero, resume el art√≠culo en un m√≠nimo de 7 lineas y un m√°ximo de 10 l√≠neas, reformulando con sin√≥nimos para evitar copiar frases literales.\n"
    "Despu√©s, valora el tono emocional impl√≠cito en el t√≠tulo del 1 (muy negativo) al 9 (muy positivo). Y 5 neutro.\n"
    "A continuaci√≥n, determina donde se producen los hechos de la publicaci√≥n o si habla en relaci√≥n a un lugar espec√≠fico. Debes deteminar la ciudad o region en el campo ciudad_region y el pais en el campo pais. Si puedes deducir el pais pero no la ciudad_region, utiliza la capital del pais. Si no se puede deducir ninguno pon la palabra indeterminado en los dos campos. El pais debe estar en formato ISO3.\n\n"
    "Devuelve el resultado √∫nicamente en formato JSON como este:\n"
    "{\n"
    "  \"resumen\": \"...\",\n"
    "  \"ciudad_region\": \"...\",\n"
    "  \"pais\": \"...\",\n"
    "  \"tono\": 5\n"
    "}\n"
    "Solo devu√©lveme el JSON. No lo envuelvas con ```json ni ning√∫n otro texto."
)


    messages = [
        {"role": "system", "content": "Eres un asistente experto en an√°lisis de prensa, resumen profesional y valoraci√≥n emocional del lenguaje."},
        {"role": "user", "content": prompt}
    ]

    respuesta = get_gpt_response(messages, temperature=0.7)

    try:
        resultado = json.loads(respuesta)
        publicacion.contenido = resultado['resumen']
        publicacion.tono = int(resultado['tono']) 
        publicacion.ciudad_region = str(resultado['ciudad_region'])
        publicacion.pais = str(resultado['pais']) 
        logging.info(f"üéØ Tono estimado: {publicacion.tono}")
        logging.info(f"‚úÖ Resumen creado: {publicacion.contenido}")
        logging.info(f"üéØ Ciudad o regi√≥n detectada: {publicacion.ciudad_region}")
        logging.info(f"üéØ Pais: {publicacion.pais}")
        return publicacion
    except (json.JSONDecodeError, KeyError, ValueError):
        raise ValueError(f"Respuesta inesperada del modelo, se esperaba JSON con claves 'resumen', 'tono', 'ciuda-region' y 'pais': {respuesta}")

MAX_TOKENS_TOTAL = 12000
TOKENS_POR_PUB = 500

#--------------------------------------------------------------


def generar_informe_impacto_temporal(publicaciones: List[dict], area_id: str, filtros: dict = None) -> BytesIO:
    if not publicaciones:
        raise ValueError("No hay publicaciones para analizar.")
    if not area_id:
        raise ValueError("Se requiere un √°rea de trabajo (area_id) para el an√°lisis.")

    publicaciones.sort(key=lambda x: x.get("fecha") or datetime.now())
    fechas = [p["fecha"] for p in publicaciones if p.get("fecha")]
    fecha_inicio = min(fechas).isoformat() if fechas else None
    fecha_fin = max(fechas).isoformat() if fechas else None

    # Obtener √°reas de impacto asociadas
    from app.mongo.mongo_area_impacto import get_areas_impacto  # importaci√≥n local
    areas_impacto = [a for a in get_areas_impacto() if str(a.area_id) == str(area_id)]

    if not areas_impacto:
        raise ValueError(f"No se encontraron √°reas de impacto para el √°rea {area_id}.")

    descripciones_impacto = "\n".join([
        f"- {a.nombre}: {a.descripcion.strip()}" for a in areas_impacto
    ])
    claves_impacto = [a.nombre.lower().strip() for a in areas_impacto]

    lote_size = MAX_TOKENS_TOTAL // TOKENS_POR_PUB
    total_lotes = math.ceil(len(publicaciones) / lote_size)
    logging.info(f"üìö Dividiendo {len(publicaciones)} publicaciones en {total_lotes} lotes para an√°lisis temporal.")

    impactos_lotes = []

    for i in range(total_lotes):
        inicio = i * lote_size
        fin = inicio + lote_size
        lote = publicaciones[inicio:fin]

        entradas = []
        for pub in lote:
            tono = pub.get("tono", "-")
            pais = pub.get("pais", "??")
            titulo = pub.get("titulo", "").strip()
            contenido = pub.get("contenido", "").strip()
            entradas.append(f"- ({tono}/10) [{pais}] {titulo}\n{contenido}")

        prompt = (
    "A partir de los siguientes titulares y contenidos de noticias ordenadas en el tiempo, redacta un an√°lisis de c√≥mo evoluciona la situaci√≥n y su impacto.\n"
    "Analiza el impacto en base a estas dimensiones definidas por el usuario:\n"
    f"{descripciones_impacto}\n\n"
    "Para cada una, resume en 1-2 l√≠neas lo que se puede deducir del conjunto de noticias.\n"
    "Indica evoluci√≥n, intensidad, y ejemplos concretos si los hay.\n"
    "Evita introducciones, explicaciones o cualquier texto adicional: este an√°lisis ser√° integrado autom√°ticamente a un informe generado por IA.\n"
    "IMPORTANTE: Devuelve exclusivamente un JSON v√°lido con esta estructura exacta:\n"
    '{\n'
    '  "impactos": {\n' +
    "".join([f'    "{clave}": "...",\n' for clave in claves_impacto]).rstrip(',\n') + '\n'
    '  }\n'
    '}\n'
    "No incluyas ning√∫n comentario, encabezado, ni c√≥digo markdown como ```json.\n"
    "Solo responde con el JSON limpio.\n"
    "Noticias:\n" + "\n\n".join(entradas[:150])
)


        messages = [
            {"role": "system", "content": "Eres un experto en an√°lisis de impacto con base en noticias."},
            {"role": "user", "content": prompt}
        ]

        try:
            data = None
            for attempt in range(2):
                respuesta = get_gpt_response(messages, temperature=0.6).strip()
                logging.debug(f"üîç Respuesta cruda (intento {attempt+1}):\n{respuesta}")
                if respuesta:
                    try:
                        data = json.loads(respuesta)
                        break
                    except json.JSONDecodeError:
                        logging.warning(f"üöß JSON inv√°lido en intento {attempt+1}")
                else:
                    logging.warning(f"üöß Respuesta vac√≠a en intento {attempt+1}")

            if not data:
                raise ValueError("No se obtuvo JSON v√°lido del modelo")

            impactos_lotes.append(data["impactos"])
            logging.info(f"‚úÖ Lote {i+1}/{total_lotes} procesado correctamente.")
        except Exception as e:
            logging.error(f"‚ùå Error en an√°lisis del lote {i+1}: {e}")

    texto_por_area = {clave: [] for clave in claves_impacto}
    for impacto in impactos_lotes:
        for clave in texto_por_area:
            texto_por_area[clave].append(impacto.get(clave, ""))

    doc = Document()
    doc.add_heading('Informe de Impacto Temporal', 0)

    if filtros:
        fuentes_map = {str(f['_id']): f for f in get_fuentes_dict()}
        concepto_id = filtros.get("concepto_interes")
        conceptos_map = {}
        if concepto_id:
            conceptos_map = {
                str(c['_id']): c
                for c in get_conceptos_collection("conceptos_interes").find({"_id": {"$in": [ObjectId(concepto_id)]}})
            }

        doc.add_heading("Filtros aplicados", level=1)
        etiquetas = {
            "concepto_interes": "Concepto de inter√©s",
            "fuente_id": "Fuente",
            "pais": "Pa√≠s",
            "tono": "Tono",
            "busqueda_palabras": "Palabras buscadas",
            "keywordsRelacionadas": "Keywords relacionadas"
        }

        if fecha_inicio and fecha_fin:
            doc.add_paragraph(f"Per√≠odo analizado: {fecha_inicio} a {fecha_fin}")

        for key, label in etiquetas.items():
            value = filtros.get(key)
            if value:
                valor_str = ""
                if key == "concepto_interes" and value in conceptos_map:
                    valor_str = conceptos_map[value].get("nombre", str(value))
                elif key == "fuente_id" and value in fuentes_map:
                    valor_str = fuentes_map[value].get("nombre", str(value))
                elif isinstance(value, list):
                    valor_str = ", ".join(map(str, value))
                else:
                    valor_str = str(value)
                doc.add_paragraph(f"{label}: {valor_str}")

    for area in claves_impacto:
        doc.add_heading(area.capitalize(), level=1)
        area_texto = "\n\n".join(texto_por_area[area]).strip()
        area_texto = resumir_parrafos_si_muchos(area_texto)
        doc.add_paragraph(area_texto)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


#-------------------------------------------

def resumir_parrafos_si_muchos(texto: str, umbral=5) -> str:
    parrafos = [p.strip() for p in texto.split("\n") if p.strip()]
    if len(parrafos) <= umbral:
        return texto

    prompt = (
        "A continuaci√≥n tienes un conjunto de p√°rrafos que resumen noticias en un √°rea espec√≠fica (econ√≥mica, social o seguridad). "
        "Por favor, genera un resumen consolidado en forma de 5 p√°rrafos como m√°ximo, conservando los puntos clave, evoluci√≥n e intensidad."
        "Evita repeticiones y enf√≥cate en los aspectos m√°s relevantes.\n\n"
        "Texto original:\n" + "\n".join(parrafos)
    )

    messages = [
        {"role": "system", "content": "Eres un experto en an√°lisis de impacto de noticias, redactas res√∫menes ejecutivos en lenguaje claro."},
        {"role": "user", "content": prompt}
    ]

    try:
        respuesta = get_gpt_response(messages, temperature=0.5).strip()
        return respuesta
    except Exception as e:
        logging.error(f"Error al resumir p√°rrafos para el √°rea: {e}")
        return "\n".join(parrafos[:5])
    

def evaluar_relacion_llm(publicacion: Publicacion, concepto: dict) -> bool:
    """
    Usa un LLM para determinar si una publicaci√≥n est√° relacionada con un concepto dado.
    El concepto debe tener al menos un campo "nombre".
    """
    if not concepto or "nombre" not in concepto:
        raise ValueError("El concepto debe tener un campo 'nombre'")

    prompt = (
        f"T√≠tulo: {publicacion.titulo}\n"
        f"Contenido: {publicacion.contenido[:1000]}...\n\n"
        f"¬øEst√° esta publicaci√≥n relacionada con el tema '{concepto['nombre']}'?\n"
        "Responde √∫nicamente con 's√≠' o 'no'. No incluyas explicaciones."
    )

    messages = [
        {"role": "system", "content": "Eres un experto en an√°lisis de contenidos y categorizaci√≥n tem√°tica."},
        {"role": "user", "content": prompt}
    ]

    try:
        respuesta = get_gpt_response(messages, temperature=0).strip().lower()
        return respuesta == "s√≠"
    except Exception as e:
        logging.error(f"‚ùå Error al consultar LLM para relaci√≥n con concepto '{concepto['nombre']}': {e}")
        return False
